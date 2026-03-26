#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown转Word文档转换脚本
用于将系统功能模块设计文档转换为Word格式
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re


def set_font(run, font_name='宋体', font_size=12, bold=False, color=None):
    """设置字体样式"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_font(run, '宋体', 16 + (2 - level) * 2, bold=True)
    return heading


def add_paragraph(doc, text, font_size=12, bold=False, indent=0):
    """添加段落"""
    para = doc.add_paragraph()
    if indent > 0:
        para.paragraph_format.left_indent = Inches(indent)
    run = para.add_run(text)
    set_font(run, '宋体', font_size, bold)
    return para


def add_table(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 添加表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_font(run, '宋体', 11, bold=True, color=RGBColor(255, 255, 255))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    # 添加数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data) if cell_data is not None else ''
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    set_font(run, '宋体', 10)


def parse_markdown(md_file):
    """解析Markdown文件"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    return lines


def convert_md_to_word(md_file, docx_file):
    """将Markdown转换为Word文档"""
    doc = Document()
    
    # 设置默认样式
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)
    
    lines = parse_markdown(md_file)
    
    i = 0
    table_headers = []
    table_rows = []
    in_table = False
    table_separator_line = False
    
    while i < len(lines):
        line = lines[i].strip()
        
        if line.startswith('#'):
            # 标题
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            add_heading(doc, text, level=level)
        
        elif line.startswith('|') and not line.startswith('|---'):
            # 表格行
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if not in_table:
                table_headers = cells
                in_table = True
            elif not table_separator_line:
                table_rows.append(cells)
        
        elif line.startswith('|---'):
            # 表格分隔线
            table_separator_line = True
        
        elif in_table and not line.startswith('|'):
            # 表格结束
            if table_headers and table_rows:
                add_table(doc, table_headers, table_rows)
            table_headers = []
            table_rows = []
            in_table = False
            table_separator_line = False
            if line:
                add_paragraph(doc, line)
        
        elif line.startswith('```'):
            # 代码块
            add_paragraph(doc, '【代码块】', font_size=10, bold=True)
        
        elif line and not line.startswith('---'):
            # 普通段落
            add_paragraph(doc, line)
        
        i += 1
    
    # 如果文档以表格结尾，添加表格
    if in_table and table_headers and table_rows:
        add_table(doc, table_headers, table_rows)
    
    # 保存文档
    doc.save(docx_file)
    print(f'✓ Word文档已成功生成：{docx_file}')


if __name__ == '__main__':
    md_file = 'SYSTEM_FUNCTION_MODULE_DESIGN_WORD.md'
    docx_file = '学生宿舍管理系统_功能模块设计.docx'
    
    try:
        convert_md_to_word(md_file, docx_file)
    except Exception as e:
        print(f'✗ 转换失败：{e}')
